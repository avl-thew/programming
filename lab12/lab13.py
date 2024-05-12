# python3 -m pip install -r requirements.txt
from abc import ABC, abstractmethod
from docx import Document
from openpyxl import Workbook
from parallelepiped import Parallelepiped
from tetrahedron import Tetrahedron
from sphere import Sphere
from kivy.app import App
from kivy.uix.button import Button
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.togglebutton import ToggleButton
from kivy.uix.behaviors.togglebutton import ToggleButtonBehavior
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
import subprocess

def save_to_doc(data, result):
    doc = Document()
    doc.add_paragraph(data)
    doc.add_paragraph(result)
    doc.save('result.docx')


def save_to_xls(data, result):
    wb = Workbook()
    ws = wb.active
    i = 2
    j = 1
    ws['A1'] = 'Вводимые данные:'
    for key, value in data.items():
        ws.cell(row=i, column=j, value=key)
        ws.cell(row=i, column=j+1, value=value)
        i = i+1
    k = 2
    p = 3

    ws['C1'] = 'Результаты:'
    for key, value in result.items():
        ws.cell(row=k, column=p, value=key)
        ws.cell(row=k, column=p+1, value=value)
        k = k+1

    wb.save('result.xlsx')

def run_lab12(instance):
    print("Запуск lab12.py...")
    path_to_lab12 = r'C:\Users\Tатьяна\OneDrive\Рабочий стол\matthew\matthew\lab12\lab12.py'
    subprocess.run(['python', path_to_lab12])
    print("lab12.py завершен.")

# class MyApp(App):
#     def build(self):
#         return Button(text='Запустить LAB 12', on_press=run_lab12)

class RadioButton(ToggleButton):
    def __init__(self, group, **kwargs):
        super(RadioButton, self).__init__(**kwargs)
        self.group = group


class RadioButtonGroup(BoxLayout):
    def __init__(self, **kwargs):
        super(RadioButtonGroup, self).__init__(**kwargs)
        self.orientation = 'vertical'
        self.add_widget(Button(text='Запустить LAB 12', on_press=run_lab12, background_color=[1, 0, 0, 1], color=[1, 1, 1, 1]))
        self.add_widget(RadioButton('shapes', text='Параллелепипед'))
        self.add_widget(RadioButton('shapes', text='Сфера'))
        self.add_widget(RadioButton('shapes', text='Тетраэдр'))
        # def build(self):
        #     return Button(text='Запустить LAB 12', on_press=run_lab12)
        
        
class Window(ABC, App):
    def __init__(self):
        self.layout = BoxLayout(orientation='vertical')
        self.fill_layout()
        super().__init__()

    @abstractmethod
    def fill_layout(self):
        pass

    @abstractmethod
    def build(self):
        pass

    @abstractmethod
    def on_button_press(self, instance):
        pass


class MainWindow(Window):
    def __init__(self, name):
        self._name = name
        super().__init__()

    def __str__(self):
        return self.name

    @property
    def name(self):
        return self._name

    @name.setter
    def set_name(self, name):
        self._name = name

    def fill_layout(self):
        self.layout.add_widget(Label(text='Выберите фигуру:'))
        self.rbg = RadioButtonGroup()
        self.layout.add_widget(self.rbg)
        self.layout.add_widget(
            Button(text='Выбрать', on_press=self.on_button_press))

    def build(self):
        return self.layout

    def on_button_press(self, instance):
        shape = None
        for child in self.rbg.children:
            if child.state == 'down':
                shape = child.text
        if not shape:
            return
        self.layout.clear_widgets()
        child = InputWindow("Окно ввода", shape, parent=self)
        child.run()


class InputWindow(Window):
    def __init__(self, name, shape, parent=None, **kwargs):
        self._parent = parent
        self._name = name
        self._shape = shape
        super().__init__(**kwargs)

    def __str__(self):
        return self._name

    @property
    def shape(self):
        return self._shape

    def fill_layout(self):
        self.layout.add_widget(Label(text='Введите размеры:'))
        if self.shape == 'Параллелепипед':
            self.layout.add_widget(Label(text='Ширина: '))
            self.layout.add_widget(TextInput(text=''))
            self.layout.add_widget(Label(text='Высота: '))
            self.layout.add_widget(TextInput(text=''))
            self.layout.add_widget(Label(text='Длина: '))
            self.layout.add_widget(TextInput(text=''))
        if self.shape == 'Сфера':
            self.layout.add_widget(Label(text='Радиус'))
            self.layout.add_widget(TextInput(text=''))
        if self.shape == 'Тетраэдр':
            self.layout.add_widget(Label(text='Сторона'))
            self.layout.add_widget(TextInput(text=''))
        self.layout.add_widget(Label(text='Плотность: '))
        self.layout.add_widget(TextInput(text=''))
        self.layout.add_widget(
            Button(text='Рассчитать', on_press=self.on_button_press))

    def build(self):
        return self.layout

    def on_button_press(self, instance):
        values = []
        for widget in self.layout.children:
            if 'textinput' in repr(widget):
                values.append(widget.text)
        if not all([value.strip().isdigit() for value in values]):
            return
        values = [int(value.strip()) for value in values]
        shapes = {
            'Параллелепипед': Parallelepiped,
            'Сфера': Sphere,
            'Тетраэдр': Tetrahedron,
            'Запустить LAB 12': run_lab12,
        }
        shape = shapes[self.shape](*values[::-1])
        self.layout.clear_widgets()
        child = ResultWindow("Окно реультатов", shape, parent=self)
        child.run()


class ResultWindow(Window):
    def __init__(self, name, shape, parent=None, **kwargs):
        self._parent = parent
        self._name = name
        self._shape = shape
        super().__init__(**kwargs)
    
    
    @property
    def shape(self):
        return self._shape

    def __str__(self):
        return self._name

    def fill_layout(self):
        self.layout.clear_widgets()
        V, S, m = self.shape.calculate()
        self.layout.add_widget(Label(text='Результаты:'))
        self.layout.add_widget(Label(text=f'Объем: {V:.2f}'))
        self.layout.add_widget(Label(text=f'Площадь: {S:.2f}'))
        self.layout.add_widget(Label(text=f'Масса: {m:.2f}'))
        self.layout.add_widget(
            Button(text='Выйти', on_press=self.on_button_press))

    def build(self):
        return self.layout

    def on_button_press(self, instance):
        attr_names = {
            'a': 'Сторона А',
            'b': 'Сторона B',
            'c': 'Сторона C',
            'R': 'Радиус',
            'side': 'Сторона',
            'density': 'Плотность',
        }
        data = {'Фигура': getattr(self.shape, 'name')}
        data = {**data, **{attr_names[str(attr)]: getattr(self.shape, attr) for attr in dir(self.shape) if not callable(
            getattr(self.shape, attr)) and not attr.startswith("__") and not attr == 'name'}}
        V, S, m = self.shape.calculate()
        result = {
            'Объем': f'{V:.2f}',
            'Площадь': f'{S:.2f}',
            'Масса': f'{m:.2f}',
        }
        save_to_doc(data, result)
        save_to_xls(data, result)
        self.layout.clear_widgets()
        self.stop()


if __name__ == '__main__':
    main_window = MainWindow("Главное окно")
    main_window.run()
    # MyApp().run()
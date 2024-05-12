import tkinter as tk
from tkinter import ttk
import tkinter.messagebox as mb
from docx import Document
from openpyxl import Workbook
from parallelepiped import Parallelepiped
from tetrahedron import Tetrahedron
from sphere import Sphere
import subprocess

def save_to_doc(data, result):
  doc = Document()
  doc.add_paragraph(data)
  doc.add_paragraph(result)
  doc.save('result.docx')

def save_to_xls(data, result):
  wb = Workbook()
  ws = wb.active
  i=2
  j=1
  ws['A1'] = 'Вводимые данные:'
  for key in data:
    ws.cell(row = i, column= j, value=key)
    ws.cell(row = i, column= j+1, value=data[key])
    i=i+1
    # ws['A2'] = key
    # ws['A3'] = data[key]
  k=1
  p=3
  
  for lov in result:
    ws.cell(row = k, column= p, value=lov)
    ws.cell(row = k, column= p+1, value=result[lov])
    k=k+1
    # ws['B2'] = result[lov]  
    # ws.cell()

  wb.save('result.xlsx')

def run_lab13():
    print("Запуск lab13.py...")
    path_to_lab13 = r'C:\Users\Tатьяна\OneDrive\Рабочий стол\matthew\matthew\lab12\lab13.py'
    subprocess.run(['python', path_to_lab13])
    print("lab13.py завершен.")
   
def open_input_window(root, shape):
    input_window = tk.Toplevel(root)
    input_window["bg"] = "#FFB6C1" # желтый цвет
    input_window.geometry("300x200") 
    
    root.withdraw()
    
    if shape == 'Параллелепипед':
        a_label = tk.Label(input_window, text='Длина:', bg="#B4EEB4")
        a_entry = tk.Entry(input_window)
        b_label = tk.Label(input_window, text='Ширина:', bg="#B4EEB4")
        b_entry = tk.Entry(input_window)
        c_label = tk.Label(input_window, text='Высота:', bg="#B4EEB4")
        c_entry = tk.Entry(input_window)
        density_label = tk.Label(input_window, text='Плотность:', bg="#B4EEB4")
        density_entry = tk.Entry(input_window)
        
        a_label.grid(row=0, column=0)
        a_entry.grid(row=0, column=1)
        b_label.grid(row=1, column=0)
        b_entry.grid(row=1, column=1)
        c_label.grid(row=2, column=0)
        c_entry.grid(row=2, column=1)
        density_label.grid(row=3, column=0)
        density_entry.grid(row=3, column=1)
        
            
        def calculate():
            a = float(a_entry.get())
            b = float(b_entry.get())
            c = float(c_entry.get())
            density = float(density_entry.get())
            p = Parallelepiped()
            V, S, m = p.calculate_parallelepiped(a, b, c, density)
            data = {'Длина': a, 'Ширина': b, 'Высота': c, 'Плотность': density}
            result = {'Объем': V, 'Площадь': S, 'Масса': m}
            save_to_doc(data, result)
            save_to_xls(data, result)
            input_window.withdraw()
            open_result_window(root, V, S, m)
            
        calc_button = tk.Button(input_window, bg="#FFE4B5")        
        calc_button = tk.Button(input_window, text='Рассчитать', command=calculate)
        calc_button.grid(row=4, column=0, columnspan=2)
        
    elif shape == 'Тетраэдр':
        a_label = tk.Label(input_window, text='Длина:', bg="#B4EEB4")
        a_entry = tk.Entry(input_window)
        density_label = tk.Label(input_window, text='Плотность:', bg="#B4EEB4")
        density_entry = tk.Entry(input_window)    
        a_label.grid(row=0, column=0)
        a_entry.grid(row=0, column=1)
        density_label.grid(row=1, column=0)
        density_entry.grid(row=1, column=1)
            
        def calculate():
            a = float(a_entry.get())
            density = float(density_entry.get())
            t = Tetrahedron()
            V, S, m = t.calculate_tetrahedron(a, density)
            data = {'Длина': a, 'Плотность': density}
            result = {'Объем': V, 'Площадь': S, 'Масса': m}
            save_to_doc(data, result)
            save_to_xls(data, result)
            input_window.withdraw()
            open_result_window(root, V, S, m)
            
        calc_button = tk.Button(input_window, bg="#FFE4B5")        
        calc_button = tk.Button(input_window, text='Рассчитать', command=calculate)
        calc_button.grid(row=2, column=0, columnspan=2)
        
    else:
        R_label = tk.Label(input_window, text='Радиус:', bg="#B4EEB4")
        R_entry = tk.Entry(input_window)
        density_label = tk.Label(input_window, text='Плотность:', bg="#B4EEB4")
        density_entry = tk.Entry(input_window)    
        R_label.grid(row=0, column=0)
        R_entry.grid(row=0, column=1)
        density_label.grid(row=1, column=0)
        density_entry.grid(row=1, column=1)
        
        def calculate():
            R = float(R_entry.get())
            density = float(density_entry.get())
            s = Sphere()
            V, S, m = s.calculate_sphere(R, density)
            data = {'Радиус': R,'Плотсноть': density}
            result = {'Объем': V, 'Площадь': S, 'Масса': m}
            save_to_doc(data, result)
            save_to_xls(data, result) 
            input_window.withdraw()
            open_result_window(root, V, S, m)
            
        calc_button = tk.Button(input_window, bg="#FFE4B5")
        calc_button = tk.Button(input_window, text='Рассчитать', command=calculate)
        calc_button.grid(row=2, column=0, columnspan=2)
        
def open_result_window(root, V, S, m):
  result_window = tk.Toplevel(root)
  result_window["bg"] = "#FFB6C1"
  result_window.geometry("300x250") 
  
  V_label = tk.Label(result_window, text=f'Объем (V): {V:.2f}')
  S_label = tk.Label(result_window, text=f'Площадь (S): {S:.2f}') 
  m_label = tk.Label(result_window, text=f'Масса (m): {m:.2f}')

  V_label.grid(row=0, column=0)
  S_label.grid(row=1, column=0)
  m_label.grid(row=2, column=0)

  ok_button = tk.Button(result_window, text='ОК', command=root.destroy)
  ok_button.grid(row=3, column=0)
  
  def open_result_window(root, V, S, m):
    def ok_button_clicked():
        V = V
        S = S 
        m = m
         
        open_result_window(root, V, S, m)
        root.destroy()   

    ok_button = tk.Button(result_window, text='ОК', command=ok_button_clicked)

def main():
  root = tk.Tk()
  root.title("LAB 12") 
  root.geometry('300x250')
  root["bg"] = '#FFB6C1'
  label = tk.Label(root, text='Выберите фигуру:', bg='#CD8C95')
  label.pack()
  
  run_lab13_button = tk.Button(root, text='Запустить LAB 13', command=run_lab13, bg='#A52A2A')
  run_lab13_button.pack()
    
  shape = tk.StringVar()

  parallelepiped_radio = tk.Radiobutton(root, text='Параллелепипед', bg='#C1FFC1', variable=shape, value='Параллелепипед')
  parallelepiped_radio.pack()

  tetrahedron_radio = tk.Radiobutton(root, text='Тетраэдр', bg='#B4EEB4', variable=shape, value='Тетраэдр')
  tetrahedron_radio.pack()  

  sphere_radio = tk.Radiobutton(root, text='Шар', bg='#9BCD9B', variable=shape, value='Шар')
  sphere_radio.pack()

  input_button = tk.Button(root, text='Ввести параметры', bg='#CD8C95', command=lambda: open_input_window(root, shape.get()))
  input_button.pack()  
  
  root.mainloop()

if __name__ == '__main__':
    main()